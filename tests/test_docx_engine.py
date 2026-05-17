"""Tests for DOCX paper export engine."""

import json
import os
import tempfile

import pytest

from docx import Document

try:
    from sophia.exporters.docx_engine import DOCXEngine
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestDOCXEngine:
    def test_create_engine(self):
        engine = DOCXEngine()
        assert engine is not None

    def test_export_basic_paper(self):
        engine = DOCXEngine()
        doc = {
            "id": "test_paper",
            "title": "Test Paper Title",
            "authors": "John Doe, Jane Smith",
            "abstract": "This is a test abstract for the paper.",
            "keywords": ["test", "paper", "docx"],
            "sections": {
                "1": {"title": "Introduction", "content": "This is the introduction."},
                "2": {"title": "Methods", "content": "We used a sample of 100 participants."},
                "3": {"title": "Results", "content": "The results were significant.\n\n[FORMULA:t:98:3.24]"},
                "4": {"title": "Discussion", "content": "These findings suggest..."},
            },
            "references": [
                "Doe, J. (2024). Test paper. Journal of Testing, 1(1), 1-10.",
            ],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            result = engine.export_paper(doc, tmp.name)
            tmp_path = tmp.name

        try:
            assert result["format"] == "docx"
            assert result["path"] == tmp_path
            assert result["sections"] == 4
            assert result["references"] == 1
            assert os.path.exists(tmp_path)
            assert os.path.getsize(tmp_path) > 1000

            # Reload and verify structure
            docx = Document(tmp_path)
            text = "\n".join(p.text for p in docx.paragraphs)
            assert "Test Paper Title" in text
            assert "John Doe, Jane Smith" in text
            assert "This is a test abstract" in text
            assert "Keywords:" in text
            assert "Introduction" in text
            assert "Methods" in text
            assert "Results" in text
            assert "Discussion" in text
            assert "References" in text
        finally:
            os.unlink(tmp_path)

    def test_export_with_inline_formula(self):
        engine = DOCXEngine()
        doc = {
            "id": "test_formula",
            "title": "Formula Test",
            "authors": "",
            "abstract": "",
            "keywords": [],
            "sections": {
                "1": {"title": "Results", "content": "The effect was significant ($t(98) = 3.24$), p < .001."},
            },
            "references": [],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            result = engine.export_paper(doc, tmp.name)
            tmp_path = tmp.name

        try:
            assert os.path.exists(tmp_path)
            docx = Document(tmp_path)
            # Check that m:oMath elements exist
            math_count = 0
            for para in docx.paragraphs:
                for child in para._p:
                    if child.tag.endswith("oMath"):
                        math_count += 1
            assert math_count >= 1, "Should have at least 1 OMML formula"
        finally:
            os.unlink(tmp_path)

    def test_export_with_apa_table(self):
        engine = DOCXEngine()
        table_json = json.dumps({
            "caption": "Table 1. Descriptive Statistics",
            "headers": ["Variable", "M", "SD", "N"],
            "rows": [["Age", "25.4", "3.2", "100"], ["Score", "78.5", "12.1", "100"]],
            "note": "Note. M = mean, SD = standard deviation.",
        })
        doc = {
            "id": "test_table",
            "title": "Table Test",
            "authors": "",
            "abstract": "",
            "keywords": [],
            "sections": {
                "1": {"title": "Results", "content": f"See the table below.\n\n[TABLE:{table_json}]"},
            },
            "references": [],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            result = engine.export_paper(doc, tmp.name)
            tmp_path = tmp.name

        try:
            assert os.path.exists(tmp_path)
            docx = Document(tmp_path)
            # Check tables exist
            assert len(docx.tables) >= 1, "Should have at least 1 table"
            table = docx.tables[0]
            assert len(table.rows) == 3  # header + 2 data rows
            assert "Variable" in table.rows[0].cells[0].text
            assert "Age" in table.rows[1].cells[0].text
        finally:
            os.unlink(tmp_path)

    def test_export_with_formula_marker(self):
        engine = DOCXEngine()
        doc = {
            "id": "test_marker",
            "title": "Formula Marker Test",
            "authors": "",
            "abstract": "",
            "keywords": [],
            "sections": {
                "1": {"title": "Results", "content": "\n[FORMULA:chi:4:12.34]\n\n[FORMULA:p:0.002]\n\n[FORMULA:ci:0.42:1.18]"},
            },
            "references": [],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            result = engine.export_paper(doc, tmp.name)
            tmp_path = tmp.name

        try:
            assert os.path.exists(tmp_path)
            docx = Document(tmp_path)
            math_count = 0
            for para in docx.paragraphs:
                for child in para._p:
                    if child.tag.endswith("oMath"):
                        math_count += 1
            assert math_count == 3, f"Should have 3 formulas, found {math_count}"
        finally:
            os.unlink(tmp_path)

    def test_export_chinese_paper(self):
        engine = DOCXEngine()
        doc = {
            "id": "test_cn",
            "title": "最低工资政策对就业影响的研究",
            "authors": "张三，李四",
            "abstract": "本文研究了最低工资政策对就业的影响。",
            "keywords": ["最低工资", "就业", "双重差分法"],
            "sections": {
                "1": {"title": "引言", "content": "最低工资政策是..."},
                "2": {"title": "文献综述", "content": "现有研究表明..."},
                "3": {"title": "研究方法", "content": "本文采用双重差分法..."},
                "4": {"title": "研究结果", "content": "结果显示政策效应显著。\n\n[FORMULA:t:98:3.24]"},
                "5": {"title": "讨论", "content": "本文发现..."},
                "6": {"title": "结论", "content": "最低工资政策..."},
            },
            "references": [
                "张三 (2024). 最低工资研究. 经济研究, 1(1), 1-20.",
            ],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            result = engine.export_paper(doc, tmp.name, citation_style="gb-t-7714-2015")
            tmp_path = tmp.name

        try:
            assert os.path.exists(tmp_path)
            docx = Document(tmp_path)
            text = "\n".join(p.text for p in docx.paragraphs)
            assert "最低工资政策对就业影响的研究" in text
            assert "张三，李四" in text
            assert "双重差分法" in text
            assert "引言" in text
            assert "研究结果" in text
        finally:
            os.unlink(tmp_path)

    def test_heading_levels(self):
        engine = DOCXEngine()
        doc = {
            "id": "test_headings",
            "title": "Heading Test",
            "authors": "",
            "abstract": "",
            "keywords": [],
            "sections": {
                "1": {"title": "Introduction", "content": "Intro text."},
                "2": {"title": "Methods", "content": "## Data\nData description.\n\n## Model\nModel spec."},
                "3": {"title": "Results", "content": "Results text."},
            },
            "references": [],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            result = engine.export_paper(doc, tmp.name)
            tmp_path = tmp.name

        try:
            docx = Document(tmp_path)
            # Check that headings are styled
            heading_styles = [p.style.name for p in docx.paragraphs if p.style.name.startswith("Heading")]
            assert len(heading_styles) >= 3, "Should have multiple headings"
        finally:
            os.unlink(tmp_path)

    def test_references_formatting(self):
        engine = DOCXEngine()
        doc = {
            "id": "test_refs",
            "title": "Ref Test",
            "authors": "",
            "abstract": "",
            "keywords": [],
            "sections": {},
            "references": [
                "Author, A. (2020). Title one. Journal One, 1(1), 1-10.",
                "Author, B. (2021). Title two. Journal Two, 2(2), 20-30.",
                "Author, C. (2022). Title three. Journal Three, 3(3), 40-50.",
            ],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            result = engine.export_paper(doc, tmp.name)
            tmp_path = tmp.name

        try:
            docx = Document(tmp_path)
            text = "\n".join(p.text for p in docx.paragraphs)
            assert "References" in text
            assert "Title one" in text
            assert "Title two" in text
            assert "Title three" in text
        finally:
            os.unlink(tmp_path)


class TestDOCXExportInterface:
    def test_export_docx_function(self):
        from sophia.exporters.docx_export import export_docx

        doc = {
            "id": "test_interface",
            "title": "Interface Test",
            "authors": "Test Author",
            "abstract": "Test abstract.",
            "keywords": [],
            "sections": {
                "1": {"title": "Section 1", "content": "Content."},
            },
            "references": [],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            result = export_docx(doc, tmp.name)
            tmp_path = tmp.name

        try:
            assert result["format"] == "docx"
            assert "error" not in result or result.get("error") is None
            assert os.path.exists(tmp_path)
        finally:
            os.unlink(tmp_path)
