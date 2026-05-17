"""Tests for OMML formula builder."""

import os
import tempfile

import pytest

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from sophia.exporters.omml_builder import (
    OMMLBuilder,
    OMMLFragment,
    GREEK_LETTERS,
    OMML_NS_DECL,
    insert_omml,
)


class TestOMMLFragment:
    def test_text_xml(self):
        frag = OMMLBuilder.text("hello")
        assert "<m:r>" in frag.xml
        assert "hello" in frag.xml

    def test_text_escapes_xml(self):
        frag = OMMLBuilder.text("a < b")
        assert "a &lt; b" in frag.xml
        frag2 = OMMLBuilder.text("x & y")
        assert "x &amp; y" in frag2.xml

    def test_add_fragments(self):
        a = OMMLBuilder.text("A")
        b = OMMLBuilder.text("B")
        c = a + b
        assert "A" in c.xml and "B" in c.xml

    def test_add_string(self):
        a = OMMLBuilder.text("A")
        c = a + "B"
        assert "A" in c.xml and "B" in c.xml


class TestSuperscript:
    def test_superscript_xml_structure(self):
        frag = OMMLBuilder.superscript("R", "2")
        assert "<m:sSup>" in frag.xml
        assert "<m:e>" in frag.xml
        assert "<m:sup>" in frag.xml
        assert "R" in frag.xml
        assert "2" in frag.xml

    def test_superscript_with_fragment(self):
        base = OMMLBuilder.greek("chi")
        frag = OMMLBuilder.superscript(base, "2")
        assert "χ" in frag.xml
        assert "2" in frag.xml


class TestSubscript:
    def test_subscript_xml_structure(self):
        frag = OMMLBuilder.subscript("x", "i")
        assert "<m:sSub>" in frag.xml
        assert "<m:e>" in frag.xml
        assert "<m:sub>" in frag.xml
        assert "x" in frag.xml
        assert "i" in frag.xml


class TestSubSup:
    def test_subsup_structure(self):
        frag = OMMLBuilder.subsup("X", "i", "j")
        assert "<m:sSubSup>" in frag.xml
        assert "<m:sub>" in frag.xml
        assert "<m:sup>" in frag.xml


class TestFraction:
    def test_fraction_structure(self):
        frag = OMMLBuilder.fraction("1", "2")
        assert "<m:f>" in frag.xml
        assert "<m:num>" in frag.xml
        assert "<m:den>" in frag.xml
        assert "1" in frag.xml
        assert "2" in frag.xml

    def test_fraction_with_fragments(self):
        num = OMMLBuilder.greek("beta")
        den = OMMLBuilder.text("SE")
        frag = OMMLBuilder.fraction(num, den)
        assert "β" in frag.xml
        assert "SE" in frag.xml


class TestSqrt:
    def test_sqrt_no_degree(self):
        frag = OMMLBuilder.sqrt("x")
        assert "<m:rad>" in frag.xml
        assert "<m:radPr>" in frag.xml
        assert "x" in frag.xml

    def test_sqrt_with_degree(self):
        frag = OMMLBuilder.sqrt("x", "3")
        assert "<m:deg>" in frag.xml
        assert "3" in frag.xml


class TestBracket:
    def test_bracket_structure(self):
        frag = OMMLBuilder.bracket("x + y")
        assert "<m:d>" in frag.xml
        assert '<m:begChr m:val="("' in frag.xml
        assert '<m:endChr m:val=")"' in frag.xml

    def test_bracket_square(self):
        frag = OMMLBuilder.bracket("x", "[", "]")
        assert '<m:begChr m:val="["' in frag.xml


class TestGreek:
    def test_greek_beta(self):
        frag = OMMLBuilder.greek("beta")
        assert "β" in frag.xml

    def test_greek_chi(self):
        frag = OMMLBuilder.greek("chi")
        assert "χ" in frag.xml

    def test_greek_all_letters(self):
        for name, char in GREEK_LETTERS.items():
            frag = OMMLBuilder.greek(name)
            assert char in frag.xml, f"Greek letter {name} should produce {char}"


class TestStatisticalExpressions:
    def test_t_stat(self):
        frag = OMMLBuilder.t_stat(98, 3.24)
        xml = frag.xml
        assert "t" in xml
        assert "98" in xml
        assert "3.24" in xml

    def test_f_stat(self):
        frag = OMMLBuilder.f_stat(3, 120, 5.67)
        xml = frag.xml
        assert "F" in xml
        assert "3" in xml
        assert "120" in xml
        assert "5.67" in xml

    def test_chi_sq(self):
        frag = OMMLBuilder.chi_sq(4, 12.34)
        xml = frag.xml
        assert "χ" in xml
        assert "4" in xml
        assert "12.34" in xml
        assert "<m:sSup>" in xml

    def test_p_value_significant(self):
        frag = OMMLBuilder.p_value(0.002)
        assert "p < .001" in frag.xml or "p =" in frag.xml

    def test_p_value_non_significant(self):
        frag = OMMLBuilder.p_value(0.042)
        assert "p = .042" in frag.xml

    def test_cohens_d(self):
        frag = OMMLBuilder.cohens_d(0.65)
        assert "Cohen's" in frag.xml
        assert "0.65" in frag.xml

    def test_r_squared(self):
        frag = OMMLBuilder.r_squared(0.847)
        assert "R" in frag.xml
        assert "2" in frag.xml
        assert "0.847" in frag.xml

    def test_r_squared_adjusted(self):
        frag = OMMLBuilder.r_squared(0.847, 0.823)
        assert "0.847" in frag.xml
        assert "0.823" in frag.xml

    def test_beta_coeff(self):
        frag = OMMLBuilder.beta_coeff("treat", 0.84, 0.12)
        assert "β" in frag.xml
        assert "treat" in frag.xml
        assert "0.840" in frag.xml or "0.84" in frag.xml
        assert "0.120" in frag.xml or "0.12" in frag.xml

    def test_ci(self):
        frag = OMMLBuilder.ci((0.42, 1.18))
        assert "95% CI" in frag.xml
        assert "0.42" in frag.xml
        assert "1.18" in frag.xml

    def test_eta_squared(self):
        frag = OMMLBuilder.eta_squared(0.142)
        assert "η" in frag.xml
        assert "0.142" in frag.xml

    def test_omega_squared(self):
        frag = OMMLBuilder.omega_squared(0.089)
        assert "ω" in frag.xml
        assert "0.089" in frag.xml

    def test_cronbach_alpha(self):
        frag = OMMLBuilder.cronbach_alpha(0.912)
        assert "Cronbach's" in frag.xml
        assert "α" in frag.xml
        assert "0.912" in frag.xml


class TestJoin:
    def test_join_multiple(self):
        frag = OMMLBuilder.join(
            OMMLBuilder.t_stat(98, 3.24),
            OMMLBuilder.p_value(0.002),
            OMMLBuilder.cohens_d(0.65),
            separator=", ",
        )
        xml = frag.xml
        assert "t" in xml
        assert "3.24" in xml
        assert "Cohen's" in xml


class TestWrapMath:
    def test_wrap_math_adds_namespace(self):
        frag = OMMLBuilder.text("x")
        wrapped = OMMLBuilder.wrap_math(frag)
        assert "<m:oMath" in wrapped
        assert "xmlns:m=" in wrapped
        assert "</m:oMath>" in wrapped


class TestInsertOMML:
    def test_insert_into_paragraph(self):
        doc = Document()
        p = doc.add_paragraph()
        frag = OMMLBuilder.t_stat(98, 3.24)
        insert_omml(p, frag)

        # Verify the paragraph contains m:oMath
        omath_found = False
        for child in p._p:
            if child.tag.endswith("oMath"):
                omath_found = True
                break
        assert omath_found, "Paragraph should contain m:oMath after insertion"

    def test_insert_roundtrip(self):
        """Save and reload document; OMML should persist."""
        doc = Document()
        p = doc.add_paragraph()
        frag = OMMLBuilder.chi_sq(4, 12.34)
        insert_omml(p, frag)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            doc2 = Document(tmp_path)
            math_count = 0
            for para in doc2.paragraphs:
                for child in para._p:
                    if child.tag.endswith("oMath"):
                        math_count += 1
            assert math_count == 1, "Should find 1 m:oMath after reload"
        finally:
            os.unlink(tmp_path)

    def test_insert_multiple_formulas(self):
        """Insert multiple different formula types into one document."""
        doc = Document()

        formulas = [
            OMMLBuilder.t_stat(98, 3.24),
            OMMLBuilder.f_stat(3, 120, 5.67),
            OMMLBuilder.chi_sq(4, 12.34),
            OMMLBuilder.p_value(0.002),
            OMMLBuilder.r_squared(0.847, 0.823),
            OMMLBuilder.beta_coeff("treat", 0.84, 0.12),
            OMMLBuilder.fraction("β", "SE"),
            OMMLBuilder.sqrt("MSE"),
        ]

        for frag in formulas:
            p = doc.add_paragraph()
            insert_omml(p, frag)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            doc2 = Document(tmp_path)
            math_count = 0
            for para in doc2.paragraphs:
                for child in para._p:
                    if child.tag.endswith("oMath"):
                        math_count += 1
            assert math_count == len(formulas), f"Should find {len(formulas)} formulas, found {math_count}"
        finally:
            os.unlink(tmp_path)


class TestComplexFormula:
    def test_full_regression_report(self):
        """Build a complex formula: β̂_treat = 0.84, SE = 0.12, t(98) = 3.24, p = .002."""
        beta = OMMLBuilder.greek("beta")
        hat = OMMLBuilder.text("̂")
        # Build beta-hat with subscript manually
        beta_hat = OMMLFragment(beta.xml + hat.xml)
        sub = OMMLBuilder.subscript(beta_hat, "treat")
        val = OMMLBuilder.text(" = 0.84, ")
        se = OMMLBuilder.text("SE = 0.12, ")
        t = OMMLBuilder.t_stat(98, 3.24)
        p = OMMLBuilder.p_value(0.002)

        full = OMMLBuilder.join(sub, val, se, t, p, separator="")
        xml = OMMLBuilder.wrap_math(full)
        assert "β" in xml
        assert "treat" in xml
        assert "0.84" in xml
        assert "SE" in xml
        assert "3.24" in xml

    def test_did_estimate(self):
        """DiD estimate with CI: β̂ = 0.84 (SE = 0.12), 95% CI [0.60, 1.07]."""
        beta = OMMLBuilder.greek("beta")
        hat = OMMLBuilder.text("̂")
        beta_hat = OMMLFragment(beta.xml + hat.xml)
        text1 = OMMLBuilder.text(" = 0.84 (SE = 0.12), ")
        ci = OMMLBuilder.ci((0.60, 1.07))

        full = OMMLBuilder.join(beta_hat, text1, ci, separator="")
        xml = full.xml
        assert "β" in xml
        assert "0.84" in xml
        assert "95% CI" in xml
        assert "0.60" in xml
        assert "1.07" in xml
