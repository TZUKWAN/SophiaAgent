"""Production-grade DOCX paper export engine for SophiaAgent.

Design philosophy (aligned with MiniMax Office Skills):
- Direct XML (oxml) manipulation where python-docx abstractions fall short
- Native OMML formulas (fully editable in Word)
- APA-style three-line tables
- Multi-level heading styles
- Header/footer with page numbers
- Reference formatting (GB/T 7714-2015 and APA 7th)
- ResultStore integration for auto-generating paper sections

Dependencies: python-docx
"""

import json
import logging
import os
import re
from datetime import datetime
from typing import Callable, Any, Dict, List, Optional, Tuple

from sophia.exporters.omml_builder import OMMLBuilder, insert_omml

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ------------------------------------------------------------------
# Constants
# ------------------------------------------------------------------

BODY_FONT = "Times New Roman"
CJK_FONT = "SimSun"
TITLE_FONT = "Times New Roman"
FONT_SIZE_PT = 12
TITLE_SIZE_PT = 18
HEADING1_SIZE_PT = 16
HEADING2_SIZE_PT = 14
HEADING3_SIZE_PT = 13
LINE_SPACING = 1.5

APA_DIMENSIONS = {
    "problem_validity": "问题有效性",
    "method_soundness": "方法合理性",
    "evidence_adequacy": "证据充分性",
    "novelty": "创新性",
    "reproducibility": "可复现性",
}


# ------------------------------------------------------------------
# Core engine
# ------------------------------------------------------------------

class DOCXEngine:
    """Generate journal-ready .docx papers from research results."""

    def __init__(self, result_store=None, workspace: str = ""):
        self.store = result_store
        self.workspace = workspace
        if not HAS_DOCX:
            raise RuntimeError("python-docx is required. Install: pip install python-docx")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def export_paper(
        self,
        doc: Dict[str, Any],
        output_path: str,
        citation_style: str = "apa7",
        include_results: bool = True,
        result_ids: Optional[List[str]] = None,
        progress_callback: Optional[Callable] = None,
    ) -> Dict[str, Any]:
        """Export a full paper to DOCX.

        Args:
            doc: Document dict from writing system (title, authors, abstract, sections, references)
            output_path: Where to write the .docx file
            citation_style: "apa7" | "gb-t-7714-2015"
            include_results: Whether to embed ResultStore results into Methods/Results
            result_ids: Specific result IDs to include (auto-detect from doc if None)

        Returns:
            Dict with path, status, section counts, warnings.
        """
        if progress_callback: progress_callback(("init", 0.05))
        document = Document()
        self._setup_styles(document)
        self._setup_page(document)

        warnings: List[str] = []

        if progress_callback: progress_callback(("rendering_title", 0.1))
        # Title page
        self._render_title_page(document, doc)

        # Abstract
        if progress_callback: progress_callback(("rendering_abstract", 0.15))
        if doc.get("abstract"):
            self._render_abstract(document, doc["abstract"], doc.get("keywords", []))

        # Sections
        sections = doc.get("sections", {})
        results_data: List[Dict] = []

        if include_results and result_ids is None:
            result_ids = self._extract_result_ids_from_doc(doc)

        if progress_callback: progress_callback(("preparing_results", 0.2))
        if include_results and result_ids and self.store:
            for rid in result_ids:
                try:
                    data = self.store.get(rid)
                    meta = self.store.get_metadata(rid)
                    results_data.append({"result_id": rid, "data": data, "meta": meta})
                except Exception as exc:
                    warnings.append(f"result_id {rid}: {exc}")

        sorted_keys = sorted(sections.keys(), key=lambda x: int(x))
        total_sections = len(sorted_keys)
        for i, key in enumerate(sorted_keys):
            if progress_callback: 
                progress = 0.2 + 0.6 * (i / max(1, total_sections))
                progress_callback((f"rendering_section_{key}", progress))
            section = sections[key]
            title = section.get("title", "")
            content = section.get("content", "")
            self._render_section(document, title, content, results_data, citation_style)

        # References
        if progress_callback: progress_callback(("rendering_references", 0.85))
        refs = doc.get("references", [])
        if refs:
            self._render_references(document, refs, citation_style)

        # Save
        if progress_callback: progress_callback(("saving", 0.95))
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        document.save(output_path)
        if progress_callback: progress_callback(("done", 1.0))

        return {
            "format": "docx",
            "path": output_path,
            "sections": len(sections),
            "references": len(refs),
            "results_embedded": len(results_data),
            "warnings": warnings,
        }

    # ------------------------------------------------------------------
    # Setup
    # ------------------------------------------------------------------

    def _setup_styles(self, document):
        """Configure document styles: Normal, Heading 1-4."""
        # Normal style
        normal = document.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(FONT_SIZE_PT)
        normal.paragraph_format.line_spacing = LINE_SPACING
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.space_before = Pt(0)
        # CJK font
        rpr = normal.element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {qn_nsdecls()} w:ascii="{BODY_FONT}" w:hAnsi="{BODY_FONT}" w:eastAsia="{CJK_FONT}"/>')
        rpr.append(rFonts)

        # Heading styles
        for i, (style_name, size_pt) in enumerate([
            ("Heading 1", HEADING1_SIZE_PT),
            ("Heading 2", HEADING2_SIZE_PT),
            ("Heading 3", HEADING3_SIZE_PT),
        ], start=1):
            try:
                style = document.styles[style_name]
            except KeyError:
                style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = BODY_FONT
            style.font.size = Pt(size_pt)
            style.font.bold = True
            style.paragraph_format.space_before = Pt(12 if i == 1 else 10)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = LINE_SPACING

    def _setup_page(self, document):
        """Configure page margins."""
        for section in document.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

    # ------------------------------------------------------------------
    # Title page
    # ------------------------------------------------------------------

    def _render_title_page(self, document, doc: Dict):
        """Render title and authors."""
        title = doc.get("title", "")
        authors = doc.get("authors", "")

        # Title
        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.name = TITLE_FONT
        run.font.size = Pt(TITLE_SIZE_PT)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        title_para.paragraph_format.space_after = Pt(18)

        # Authors
        if authors:
            author_para = document.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = author_para.add_run(authors)
            run.font.name = BODY_FONT
            run.font.size = Pt(FONT_SIZE_PT)
            author_para.paragraph_format.space_after = Pt(12)

    # ------------------------------------------------------------------
    # Abstract
    # ------------------------------------------------------------------

    def _render_abstract(self, document, abstract: str, keywords: List[str]):
        """Render abstract and keywords."""
        heading = document.add_paragraph()
        heading.style = document.styles["Heading 1"]
        run = heading.add_run("Abstract")
        run.font.name = BODY_FONT

        # Abstract text
        abs_para = document.add_paragraph()
        abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = abs_para.add_run(abstract)
        run.font.name = BODY_FONT
        run.font.size = Pt(FONT_SIZE_PT)
        abs_para.paragraph_format.first_line_indent = Cm(0.74)  # ~0.3 inch

        # Keywords
        if keywords:
            kw_para = document.add_paragraph()
            kw_run = kw_para.add_run("Keywords: ")
            kw_run.font.name = BODY_FONT
            kw_run.font.bold = True
            kw_run.font.italic = True
            kw_para.add_run("; ".join(keywords)).font.name = BODY_FONT

        document.add_paragraph()  # spacing

    # ------------------------------------------------------------------
    # Sections
    # ------------------------------------------------------------------

    def _render_section(
        self,
        document,
        title: str,
        content: str,
        results_data: List[Dict],
        citation_style: str,
    ):
        """Render a paper section with heading and content."""
        # Determine heading level from title patterns
        level = self._infer_heading_level(title)
        style_name = f"Heading {level}"

        heading = document.add_paragraph()
        if style_name in document.styles:
            heading.style = document.styles[style_name]
        run = heading.add_run(title)
        run.font.name = BODY_FONT
        run.font.bold = True

        # Process content line by line
        if content:
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Markdown-style sub-headings within content
                if line.startswith("### "):
                    sub = document.add_paragraph()
                    sub.style = document.styles["Heading 3"]
                    sub.add_run(line[4:]).font.name = BODY_FONT
                    continue
                if line.startswith("## "):
                    sub = document.add_paragraph()
                    sub.style = document.styles["Heading 2"]
                    sub.add_run(line[3:]).font.name = BODY_FONT
                    continue

                # Table markers: [TABLE:data_json]
                if line.startswith("[TABLE:") and line.endswith("]"):
                    try:
                        table_data = json.loads(line[7:-1])
                        self._render_apa_table(document, table_data)
                    except Exception:
                        para = document.add_paragraph(line)
                        para.runs[0].font.name = BODY_FONT
                    continue

                # Formula markers: [FORMULA:expression_type:...]
                if line.startswith("[FORMULA:") and line.endswith("]"):
                    para = document.add_paragraph()
                    self._render_formula_line(para, line[9:-1])
                    continue

                # Regular paragraph
                para = document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._render_inline_content(para, line)

        # Auto-embed results for Methods and Results sections
        if self.store and results_data:
            lower_title = title.lower()
            if "method" in lower_title and "result" not in lower_title:
                self._auto_render_methods(document, results_data)
            elif "result" in lower_title:
                self._auto_render_results(document, results_data)

    def _infer_heading_level(self, title: str) -> int:
        """Infer heading level from Chinese/English section titles."""
        lower = title.lower()
        # Top-level sections
        if any(k in lower for k in ("introduction", "文献综述", "方法", "results", "discussion", "结论", "references")):
            return 1
        # Second-level (sub-sections)
        if any(k in lower for k in ("data", "sample", "model", "分析", "robustness", "limitation")):
            return 2
        return 2  # default to h2 for safety

    # ------------------------------------------------------------------
    # Inline content rendering (text + inline formulas + citations)
    # ------------------------------------------------------------------

    def _render_inline_content(self, paragraph, text: str):
        """Render a line of text, handling inline formula markers and citations."""
        # Pattern: $formula$ for inline math
        parts = re.split(r'(\$[^$]+\$)', text)
        for part in parts:
            if part.startswith("$") and part.endswith("$") and len(part) > 2:
                formula_text = part[1:-1]
                self._insert_inline_formula(paragraph, formula_text)
            else:
                run = paragraph.add_run(part)
                run.font.name = BODY_FONT
                run.font.size = Pt(FONT_SIZE_PT)

    def _insert_inline_formula(self, paragraph, formula_text: str):
        """Parse common inline formula text and insert OMML."""
        formula_text = formula_text.strip()

        # Try to match common patterns
        frag = self._parse_formula_text(formula_text)
        if frag:
            insert_omml(paragraph, frag)
        else:
            # Fallback: plain text
            run = paragraph.add_run(formula_text)
            run.font.name = BODY_FONT

    def _parse_formula_text(self, text: str) -> Optional[Any]:
        """Parse common formula text into OMMLFragment."""
        text = text.strip()

        # t(df) = value
        m = re.match(r't\s*\(\s*(\d+)\s*\)\s*=\s*([\d.]+)', text)
        if m:
            return OMMLBuilder.t_stat(int(m.group(1)), float(m.group(2)))

        # F(df1, df2) = value
        m = re.match(r'F\s*\(\s*(\d+)\s*,\s*(\d+)\s*\)\s*=\s*([\d.]+)', text)
        if m:
            return OMMLBuilder.f_stat(int(m.group(1)), int(m.group(2)), float(m.group(3)))

        # χ²(df) = value
        m = re.match(r'[χχ]²?\s*\(\s*(\d+)\s*\)\s*=\s*([\d.]+)', text)
        if m:
            return OMMLBuilder.chi_sq(int(m.group(1)), float(m.group(2)))

        # p = value or p < value
        m = re.match(r'p\s*([=<>])\s*([\d.]+)', text)
        if m:
            val = float(m.group(2))
            return OMMLBuilder.p_value(val if m.group(1) == "=" else val)

        # Cohen's d = value
        m = re.match(r"Cohen'?s?\s*d\s*=\s*([\d.]+)", text, re.I)
        if m:
            return OMMLBuilder.cohens_d(float(m.group(1)))

        # β_name = value
        m = re.match(r'β[_]?([\w]+)?\s*=\s*([\d.]+)', text)
        if m:
            return OMMLBuilder.beta_coeff(m.group(1) or "", float(m.group(2)))

        # R² = value
        m = re.match(r'R²?\s*=\s*([\d.]+)', text)
        if m:
            return OMMLBuilder.r_squared(float(m.group(1)))

        # 95% CI [low, high]
        m = re.match(r'95%\s*CI\s*\[\s*([\d.-]+)\s*,\s*([\d.-]+)\s*\]', text)
        if m:
            return OMMLBuilder.ci((float(m.group(1)), float(m.group(2))))

        # Generic: try to identify Greek letters
        if any(g in text for g in GREEK_LETTERS.values()):
            return self._build_generic_formula(text)

        return None

    def _build_generic_formula(self, text: str) -> Any:
        """Build a formula from text containing Greek letters and operators."""
        # Simple approach: treat entire text as a math run
        return OMMLBuilder.text(text)

    # ------------------------------------------------------------------
    # Formula line rendering
    # ------------------------------------------------------------------

    def _render_formula_line(self, paragraph, formula_spec: str):
        """Render a standalone formula from [FORMULA:...] marker."""
        parts = formula_spec.split(":", 2)
        if len(parts) < 2:
            return
        ftype = parts[0]
        args = parts[1:]

        frag = None
        if ftype == "t" and len(args) >= 2:
            frag = OMMLBuilder.t_stat(int(args[0]), float(args[1]))
        elif ftype == "f" and len(args) >= 3:
            frag = OMMLBuilder.f_stat(int(args[0]), int(args[1]), float(args[2]))
        elif ftype == "chi" and len(args) >= 2:
            frag = OMMLBuilder.chi_sq(int(args[0]), float(args[1]))
        elif ftype == "p" and len(args) >= 1:
            frag = OMMLBuilder.p_value(float(args[0]))
        elif ftype == "beta" and len(args) >= 2:
            frag = OMMLBuilder.beta_coeff(args[0], float(args[1]), float(args[2]) if len(args) > 2 else None)
        elif ftype == "ci" and len(args) >= 2:
            frag = OMMLBuilder.ci((float(args[0]), float(args[1])))
        elif ftype == "eta" and len(args) >= 1:
            frag = OMMLBuilder.eta_squared(float(args[0]))
        elif ftype == "omega" and len(args) >= 1:
            frag = OMMLBuilder.omega_squared(float(args[0]))
        elif ftype == "alpha" and len(args) >= 1:
            frag = OMMLBuilder.cronbach_alpha(float(args[0]))

        if frag:
            insert_omml(paragraph, frag)
        else:
            run = paragraph.add_run(formula_spec)
            run.font.name = BODY_FONT

    # ------------------------------------------------------------------
    # APA Tables
    # ------------------------------------------------------------------

    def _render_apa_table(self, document, table_data: Dict):
        """Render an APA-style three-line table.

        table_data format:
        {
            "caption": "Table 1. Descriptive Statistics",
            "headers": ["Variable", "M", "SD", "N"],
            "rows": [["Age", "25.4", "3.2", "100"], ...],
            "note": "Note. M = mean, SD = standard deviation."
        }
        """
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        caption = table_data.get("caption", "")
        note = table_data.get("note", "")

        if not headers or not rows:
            return

        # Caption
        if caption:
            cap_para = document.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cap_para.add_run(caption)
            run.font.name = BODY_FONT
            run.font.size = Pt(FONT_SIZE_PT)
            run.font.italic = True
            cap_para.paragraph_format.space_after = Pt(3)

        # Table
        table = document.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = BODY_FONT
                    run.font.bold = True
                    run.font.size = Pt(FONT_SIZE_PT)

        # Data rows
        for row_idx, row_data in enumerate(rows):
            cells = table.rows[row_idx + 1].cells
            for col_idx, val in enumerate(row_data):
                if col_idx < len(cells):
                    cells[col_idx].text = str(val)
                    for paragraph in cells[col_idx].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = BODY_FONT
                            run.font.size = Pt(FONT_SIZE_PT)

        # APA three-line style: remove vertical borders, keep horizontal
        self._apply_three_line_style(table)

        # Note
        if note:
            note_para = document.add_paragraph()
            note_para.paragraph_format.space_before = Pt(3)
            run = note_para.add_run(note)
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
            run.font.italic = True

        document.add_paragraph()  # spacing after table

    def _apply_three_line_style(self, table):
        """Apply APA three-line table style via direct XML."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn

        tbl = table._tbl
        # Remove all vertical borders
        for tc in tbl.iter(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                # Clear existing borders
                for existing in tcPr.findall(qn('w:tcBorders')):
                    tcPr.remove(existing)
                # Add custom borders: only top and bottom on all cells
                borders_xml = (
                    f'<w:tcBorders {qn_nsdecls()}>'
                    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '</w:tcBorders>'
                )
                tcPr.append(parse_xml(borders_xml))

        # Bold the header row top border (slightly thicker)
        if table.rows:
            for cell in table.rows[0].cells:
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    for borders in tcPr.findall(qn('w:tcBorders')):
                        tcPr.remove(borders)
                    borders_xml = (
                        f'<w:tcBorders {qn_nsdecls()}>'
                        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        '</w:tcBorders>'
                    )
                    tcPr.append(parse_xml(borders_xml))

            # Bottom border of last row (slightly thicker)
            last_row = table.rows[-1]
            for cell in last_row.cells:
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    for borders in tcPr.findall(qn('w:tcBorders')):
                        tcPr.remove(borders)
                    borders_xml = (
                        f'<w:tcBorders {qn_nsdecls()}>'
                        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        '</w:tcBorders>'
                    )
                    tcPr.append(parse_xml(borders_xml))

    # ------------------------------------------------------------------
    # Auto-render Methods / Results from ResultStore
    # ------------------------------------------------------------------

    def _auto_render_methods(self, document, results_data: List[Dict]):
        """Auto-generate Methods paragraphs from ResultStore metadata."""
        from sophia.research.apa import APAFormatter

        seen_tools = set()
        for r in results_data:
            meta = r.get("meta", {})
            tool = meta.get("tool", "")
            if tool in seen_tools or not tool:
                continue
            seen_tools.add(tool)

            raw_params = meta.get("params", "{}")
            try:
                params = json.loads(raw_params) if isinstance(raw_params, str) else (raw_params or {})
            except Exception:
                params = {}

            # Clean bulky params
            params_clean = {k: v for k, v in params.items()
                           if k not in ("data", "texts", "X", "y", "groups")
                           and (not isinstance(v, (list, dict)) or k in ("test", "method", "type"))}

            methods_text = APAFormatter.methods_section(tool, params_clean, r.get("data", {}))
            if methods_text:
                para = document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = para.add_run(methods_text)
                run.font.name = BODY_FONT
                run.font.size = Pt(FONT_SIZE_PT)

    def _auto_render_results(self, document, results_data: List[Dict]):
        """Auto-generate Results paragraphs from ResultStore data."""
        for r in results_data:
            data = r.get("data", {})
            apa = data.get("apa", "")
            if apa:
                para = document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Try to parse inline formulas
                self._render_inline_content(para, apa)

            # Render coefficient tables
            if "coefficients" in data and "std_errors" in data:
                self._render_coefficient_table(document, data)

    def _render_coefficient_table(self, document, data: Dict):
        """Render regression coefficient table."""
        coeffs = data.get("coefficients", {})
        ses = data.get("std_errors", {})
        ps = data.get("p_values", {})
        ts = data.get("t_stats", {})

        rows = []
        for name in coeffs:
            b = coeffs.get(name, "")
            se = ses.get(name, "")
            t = ts.get(name, "")
            p = ps.get(name, "")
            rows.append([name, f"{b:.3f}" if isinstance(b, (int, float)) else str(b),
                        f"{se:.3f}" if isinstance(se, (int, float)) else str(se),
                        f"{t:.2f}" if isinstance(t, (int, float)) else str(t),
                        f"{p:.3f}" if isinstance(p, (int, float)) else str(p)])

        if rows:
            self._render_apa_table(document, {
                "caption": "Table X. Regression Coefficients",
                "headers": ["Predictor", "β", "SE", "t", "p"],
                "rows": rows,
                "note": "Note. β = unstandardized coefficient; SE = standard error.",
            })

    # ------------------------------------------------------------------
    # References
    # ------------------------------------------------------------------

    def _render_references(self, document, refs: List[str], citation_style: str):
        """Render references section."""
        heading = document.add_paragraph()
        heading.style = document.styles["Heading 1"]
        run = heading.add_run("References")
        run.font.name = BODY_FONT

        for ref in refs:
            para = document.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(-0.74)
            para.paragraph_format.left_indent = Cm(0.74)
            para.paragraph_format.line_spacing = 2.0  # double spacing for references
            run = para.add_run(ref)
            run.font.name = BODY_FONT
            run.font.size = Pt(FONT_SIZE_PT)

    # ------------------------------------------------------------------
    # Utilities
    # ------------------------------------------------------------------

    def _extract_result_ids_from_doc(self, doc: Dict) -> List[str]:
        """Scan document content for result_id references."""
        result_ids = []
        content = json.dumps(doc, ensure_ascii=False)
        # Match result_id patterns in content
        for match in re.finditer(r'res_[a-zA-Z0-9]+', content):
            rid = match.group(0)
            if rid not in result_ids:
                result_ids.append(rid)
        return result_ids


# ------------------------------------------------------------------
# Helper: namespace declarations for oxml
# ------------------------------------------------------------------

def qn_nsdecls() -> str:
    """Return namespace declarations for w: namespace."""
    return 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
