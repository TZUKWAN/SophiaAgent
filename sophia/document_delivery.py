"""Final document materialization for agent answers."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Literal, Optional

from sophia.paper_quality import is_paper_generation_request

OutputFormat = Literal["docx", "markdown", "pdf", "latex"]


def requested_output_format(user_message: str) -> Optional[OutputFormat]:
    text = user_message.lower()
    if any(term in text for term in ["docx", ".docx", "word", "word文档", "word 文件"]):
        return "docx"
    if any(term in text for term in ["pdf", ".pdf"]):
        return "pdf"
    if any(term in text for term in ["latex", ".tex", "tex文件", "tex 文档"]):
        return "latex"
    if any(term in text for term in ["markdown", ".md", "md文档", "md 文件"]):
        return "markdown"
    return None


def save_generated_docx(workspace: str, user_message: str, content: str) -> Optional[str]:
    if not content.strip() or not is_paper_generation_request(user_message):
        return None

    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        return None

    title = _extract_title(user_message) or _extract_title(content) or "sophia_generated_paper"
    output_dir = Path(workspace).expanduser().resolve() / ".sophia" / "generated_documents"
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{_safe_filename(title)}.docx"

    doc = Document()
    styles = doc.styles
    if "Normal" in styles:
        styles["Normal"].font.name = "宋体"
        styles["Normal"].font.size = Pt(12)

    _write_markdownish_content(doc, content)
    doc.save(path)
    return str(path)


def _write_markdownish_content(doc, content: str) -> None:
    lines = content.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i].rstrip()

        if line.strip().startswith("```"):
            if in_code:
                _add_preformatted(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        table_block = _consume_markdown_table(lines, i)
        if table_block:
            _add_table(doc, table_block)
            i += len(table_block)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            doc.add_heading(_strip_inline_markdown(heading.group(2)), level=level)
        elif line.strip() in {"---", "***"}:
            doc.add_paragraph("")
        elif line.strip():
            doc.add_paragraph(_strip_inline_markdown(line))
        i += 1

    if code_lines:
        _add_preformatted(doc, "\n".join(code_lines))


def _consume_markdown_table(lines: list[str], start: int) -> list[str]:
    if start + 1 >= len(lines):
        return []
    if "|" not in lines[start] or "|" not in lines[start + 1]:
        return []
    if not re.match(r"^\s*\|?[\s:\-\|]+\|?\s*$", lines[start + 1]):
        return []
    block = [lines[start], lines[start + 1]]
    idx = start + 2
    while idx < len(lines) and "|" in lines[idx] and lines[idx].strip():
        block.append(lines[idx])
        idx += 1
    return block


def _add_table(doc, block: list[str]) -> None:
    rows = [_split_table_row(line) for line in block if line.strip()]
    if len(rows) < 2:
        return
    rows = [rows[0]] + rows[2:]
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for row_values in rows:
        cells = table.add_row().cells
        for idx in range(width):
            cells[idx].text = row_values[idx] if idx < len(row_values) else ""


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [_strip_inline_markdown(cell.strip()) for cell in stripped.split("|")]


def _add_preformatted(doc, text: str) -> None:
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.*?)\1", r"\2", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _extract_title(*texts: str) -> str:
    for text in texts:
        for line in text.splitlines():
            cleaned = line.strip().lstrip("#").strip()
            if 4 <= len(cleaned) <= 80 and not cleaned.startswith(("-", "|")):
                return cleaned
    return ""


def _safe_filename(text: str) -> str:
    cleaned = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", text, flags=re.UNICODE)
    cleaned = cleaned.strip("._")
    return (cleaned or "sophia_generated_paper")[:80]
